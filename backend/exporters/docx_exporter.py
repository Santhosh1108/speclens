"""
Generates a presentation-ready Word document (.docx) from a ProductModel
and, optionally, a critique dict. Produces a title page, styled headings,
tables for requirements/metrics/risks, and a MoSCoW-colored priority tag.
"""

import io
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schemas.product_model import ProductModel

BRAND_COLOR = RGBColor(0x2A, 0x2E, 0x45)
ACCENT_COLOR = RGBColor(0x4F, 0x6B, 0xFF)
MUTED_COLOR = RGBColor(0x6B, 0x70, 0x80)

PRIORITY_COLORS = {
    "must": "C0392B",
    "should": "D68910",
    "could": "2E86C1",
    "wont": "7F8C8D",
}

SEVERITY_COLORS = {
    "high": "C0392B",
    "medium": "D68910",
    "low": "2E86C1",
}


def _shade_cell(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_COLOR
    return h


def _add_bullets(doc, items, empty_text="Not specified"):
    if not items:
        doc.add_paragraph(empty_text, style="Intense Quote")
        return
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        _set_cell_text(hdr_cells[i], header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr_cells[i], "2A2E45")

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            _set_cell_text(row_cells[i], str(value))

    return table


def build_prd_docx(
    product: ProductModel,
    prd_markdown: str = "",
    critique: dict | None = None,
) -> io.BytesIO:
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---------------- Title page ----------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(product.product or "Product Requirements Document")
    title_run.font.size = Pt(30)
    title_run.font.bold = True
    title_run.font.color.rgb = BRAND_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("Product Requirements Document")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = ACCENT_COLOR

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(date.today().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = MUTED_COLOR

    doc.add_page_break()

    # ---------------- Problem ----------------
    _add_heading(doc, "Problem", level=1)
    doc.add_paragraph(product.problem or "Not specified")

    # ---------------- Target users / context / goals ----------------
    _add_heading(doc, "Target Users", level=1)
    _add_bullets(doc, product.users)

    _add_heading(doc, "Current Context", level=1)
    _add_bullets(doc, product.current_context)

    _add_heading(doc, "Goals", level=1)
    _add_bullets(doc, product.goals)

    # ---------------- MVP scope ----------------
    _add_heading(doc, "MVP Scope", level=1)
    _add_bullets(doc, product.mvp_scope)

    _add_heading(doc, "Out of Scope", level=1)
    _add_bullets(doc, product.out_of_scope)

    # ---------------- Requirements table ----------------
    _add_heading(doc, "Requirements", level=1)
    if product.requirements:
        rows = [(r.priority.upper(), r.description, r.type) for r in product.requirements]
        _add_table(doc, ["Priority", "Description", "Type"], rows)
    else:
        doc.add_paragraph("Not specified", style="Intense Quote")

    _add_heading(doc, "Non-Functional Requirements", level=1)
    _add_bullets(doc, product.non_functional_requirements)

    # ---------------- User stories ----------------
    _add_heading(doc, "User Stories", level=1)
    if product.user_stories:
        for story in product.user_stories:
            doc.add_paragraph(
                f"As a {story.actor}, I want to {story.action}, so that {story.goal}.",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("Not specified", style="Intense Quote")

    # ---------------- Acceptance criteria ----------------
    _add_heading(doc, "Acceptance Criteria", level=1)
    _add_bullets(doc, [c.description for c in product.acceptance_criteria])

    # ---------------- Edge cases ----------------
    _add_heading(doc, "Edge Cases", level=1)
    _add_bullets(doc, product.edge_cases)

    # ---------------- Success metrics ----------------
    _add_heading(doc, "Success Metrics", level=1)
    if product.success_metrics:
        rows = [(m.name, m.target) for m in product.success_metrics]
        _add_table(doc, ["Metric", "Target"], rows)
    else:
        doc.add_paragraph("Not specified", style="Intense Quote")

    # ---------------- Risks ----------------
    _add_heading(doc, "Risks & Mitigations", level=1)
    if product.risks:
        rows = [(r.severity.upper(), r.description, r.mitigation) for r in product.risks]
        _add_table(doc, ["Severity", "Risk", "Mitigation"], rows)
    else:
        doc.add_paragraph("Not specified", style="Intense Quote")

    # ---------------- Roadmap ----------------
    _add_heading(doc, "Roadmap", level=1)
    if product.roadmap:
        rows = [(p.name, p.description) for p in product.roadmap]
        _add_table(doc, ["Phase", "Description"], rows)
    else:
        doc.add_paragraph("Not specified", style="Intense Quote")

    # ---------------- Open questions ----------------
    _add_heading(doc, "Open Questions", level=1)
    _add_bullets(doc, product.open_questions, empty_text="None")

    # ---------------- Critique appendix ----------------
    if critique:
        doc.add_page_break()
        _add_heading(doc, "Appendix: PRD Critique", level=1)

        score_p = doc.add_paragraph()
        score_run = score_p.add_run(f"Overall Score: {critique.get('overall_score', 'N/A')}/100")
        score_run.bold = True
        score_run.font.size = Pt(13)
        score_run.font.color.rgb = ACCENT_COLOR

        if critique.get("summary"):
            doc.add_paragraph(critique["summary"])

        if critique.get("strengths"):
            _add_heading(doc, "Strengths", level=2)
            _add_bullets(doc, critique["strengths"])

        if critique.get("issues"):
            _add_heading(doc, "Issues", level=2)
            rows = [
                (i.get("severity", "").upper(), i.get("category", ""), i.get("issue", ""), i.get("suggestion", ""))
                for i in critique["issues"]
            ]
            _add_table(doc, ["Severity", "Category", "Issue", "Suggestion"], rows)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
